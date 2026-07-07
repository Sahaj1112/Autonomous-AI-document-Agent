import os
import re
import uuid
import logging
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from application.ports.document_port import DocumentPort
from domain.exceptions.domain_exceptions import DocumentGenerationError

logger = logging.getLogger(__name__)


class DocxDocumentAdapter(DocumentPort):
    """Concrete implementation of DocumentPort using python-docx to write and format .docx files."""

    def __init__(self) -> None:
        pass

    def generate_document(self, content: str, document_type: str) -> bytes:
        """Generates a professional business document."""
        try:
            doc = Document()

            # --- Layout Margins ---
            section = doc.sections[0]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

            # --- Primary Styling ---
            style = doc.styles["Normal"]
            font = style.font
            font.name = "Calibri"
            font.size = Pt(11)
            font.color.rgb = RGBColor(0x33, 0x33, 0x33)

            self._configure_headings(doc)

            # --- Cover / Title Page ---
            self._add_title_page(doc, document_type)

            # --- Content Insertion ---
            doc.add_page_break()

            # Split consolidated content by the section separator '---'
            # (Note: we split on boundaries that can optionally contain newlines or spaces)
            raw_sections = re.split(r"\n*---\n*", content)
            sections_added = 0

            for section_block in raw_sections:
                section_block = section_block.strip()
                if not section_block:
                    continue

                # Split description heading from section content
                # The first paragraph represents the section description/title
                parts = section_block.split("\n\n", 1)
                if not parts:
                    continue

                heading = parts[0].strip()
                body = parts[1].strip() if len(parts) > 1 else ""

                self._add_section(doc, heading, body)
                sections_added += 1

            # --- Footer ---
            self._add_footer(doc)

            # --- Save to Bytes ---
            buffer = io.BytesIO()
            doc.save(buffer)
            logger.info("Document compiled to binary successfully.")

            return buffer.getvalue()

        except Exception as e:
            logger.error("Failed to generate Word document: %s", str(e), exc_info=True)
            raise DocumentGenerationError(f"Word document generation failed: {str(e)}") from e

    def _configure_headings(self, doc: Document) -> None:
        # Title/Heading 1 Style
        h1_style = doc.styles["Heading 1"]
        h1_font = h1_style.font
        h1_font.name = "Calibri"
        h1_font.size = Pt(18)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)  # Navy Blue

        # Subtitle/Heading 2 Style
        h2_style = doc.styles["Heading 2"]
        h2_font = h2_style.font
        h2_font.name = "Calibri"
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(0x2D, 0x5F, 0x9A)  # Slate Blue

    def _add_title_page(self, doc: Document, document_type: str) -> None:
        # Add padding before title
        for _ in range(4):
            doc.add_paragraph("")

        # Document Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run(document_type)
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

        # Divider
        doc.add_paragraph("")
        p_div = doc.add_paragraph()
        p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_div.add_run("─" * 40)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Meta/Subtitle
        for _ in range(4):
            doc.add_paragraph("")

        meta_lines = [
            f"Date generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
            "Author: Autonomous AI Document Agent",
            "Classification: Confidential Business Use Only",
        ]

        for line in meta_lines:
            p_meta = doc.add_paragraph()
            p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_meta.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.font.italic = True

    def _add_section(self, doc: Document, heading: str, content: str) -> None:
        # Heading 1 (Clean up any task numbers)
        clean_heading = re.sub(r"^\d+[\.\)]\s*", "", heading).strip()
        doc.add_heading(clean_heading, level=1)

        # Paragraph parsing
        paragraphs = content.split("\n")
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # List formatters
            if para.startswith(("- ", "* ", "• ")):
                doc.add_paragraph(para[2:].strip(), style="List Bullet")
            elif re.match(r"^\d+[\.\)]\s", para):
                clean_list_item = re.sub(r"^\d+[\.\)]\s*", "", para)
                doc.add_paragraph(clean_list_item.strip(), style="List Number")
            elif para.endswith(":") and len(para) < 100:
                # Sub-heading representation
                doc.add_heading(para.rstrip(":"), level=2)
            else:
                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(6)

    def _add_footer(self, doc: Document) -> None:
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            f"Confidential | Generated by AI Document Agent | {datetime.now().strftime('%Y-%m-%d')}"
        )
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    def _generate_filename(self, document_type: str) -> str:
        slug = re.sub(r"[^a-zA-Z0-9]+", "_", document_type).strip("_").lower()
        uuid_suffix = uuid.uuid4().hex[:8]
        return f"{slug}_{uuid_suffix}.docx"
